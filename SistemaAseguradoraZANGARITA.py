
from pydoc import Doc
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_BREAK
import openpyxl
import re
from time import strftime, strptime
from CapaNegocio.Autos import *
from CapaNegocio.Personas import *
from CapaNegocio.Informes import *
from CapaNegocio.Siniestro import *
from datetime import date



import sys
from PyQt5 import QtWidgets, uic

from CapaNegocio.Siniestro import Siniestro

class iniciar:
    def __init__(self) -> None:
        app=QtWidgets.QApplication([])
        self.ventana=uic.loadUi(r"C:\Users\Zangara\Documents\1-CURSO AX\8460 - TÉCNICAS AVANZADAS DE PROGRAMACIÓN\BBDD\.vscode\LaSegura\CapaVista\VentanaPrincipal.ui")
        self.ventana.show()

        ############conectamos menú con funciones ######################
        self.ventana.actionSalir.triggered.connect(self.accionSalir)
        self.ventana.actionAutos.triggered.connect(self.accionCRUDauto)
        self.ventana.actionPersonas.triggered.connect(self.accionCRUDPersona)
        self.ventana.actionSiniestros.triggered.connect(self.accionSiniestros)
        self.ventana.actionAccidentes.triggered.connect(self.accionAccidentes)

        app.exec()

        ##################definimos funciones#############################
    def accionSalir(self):
        sys.exit()

    # def accionCerrar(self):
    #     self.crudautos.close()

    def accionCRUDauto(self):
        self.crudautos=uic.loadUi(r"C:\Users\Zangara\Documents\1-CURSO AX\8460 - TÉCNICAS AVANZADAS DE PROGRAMACIÓN\BBDD\.vscode\LaSegura\CapaVista\CRUDautos.ui")
        self.crudautos.show()

   
    ##############obtener valores del los input#######################

        # self.matricula= self.crudautos.txtMatricula.text()
        # self.anio=self.crudautos.txtAnio.text()
        # self.modelo=self.crudautos.txtModelo.text()

    #############conectamos los botones#############
        self.crudautos.btnAgregar.clicked.connect(self.click_btnAgregar)
        self.crudautos.btnEliminar.clicked.connect(self.click_btnEliminar)
        self.crudautos.btnBuscar.clicked.connect(self.click_btnBuscar)
        self.crudautos.btnCerrar.clicked.connect(self.crudautos.close)
        self.crudautos.btnModificar.clicked.connect(self.click_btnModificar)
    def accionCRUDPersona(self):
        self.crudPersonas=uic.loadUi(r"C:\Users\Zangara\Documents\1-CURSO AX\8460 - TÉCNICAS AVANZADAS DE PROGRAMACIÓN\BBDD\.vscode\LaSegura\CapaVista\CRUDpersonas.ui")
        self.crudPersonas.show()

        self.crudPersonas.btnAgregar.clicked.connect(self.click_btnAgregarPersona)
        self.crudPersonas.btnEliminar.clicked.connect(self.click_btnEliminarPersona)
        self.crudPersonas.btnBuscar.clicked.connect(self.click_btnBuscarPersona)
        self.crudPersonas.btnCerrar.clicked.connect(self.accionCerrarPersona)
        self.crudPersonas.btnModificar.clicked.connect(self.click_btnModificarPersona)


    def accionSiniestros(self):
        self.Siniestros=uic.loadUi(r"C:\Users\Zangara\Documents\1-CURSO AX\8460 - TÉCNICAS AVANZADAS DE PROGRAMACIÓN\BBDD\.vscode\LaSegura\CapaVista\VentanaAseguradora.ui")
        self.Siniestros.show()
        self.Siniestros.btnExcel.clicked.connect(self.click_btnExcel)
        self.Siniestros.btnWord.clicked.connect(self.click_btnWord)

    def click_btnWord(sefl):
        doc= docx.Document()
        
        doc.add_heading('Listado General de Personas', 0)
        tpersonas=Persona()
        personas=tpersonas.getTodos()
       
        doc.add_paragraph('CARNET            NOMBRE Y APELLIDO     DOMICILIO')
        for persona in personas:
            doc.add_paragraph('{}            {}     {}'.format(persona[1],persona[2], persona[3] ))
        doc.save('listadoPersonas.docx')
    def click_btnExcel(self):

        libro=openpyxl.Workbook() 
    #fijar la hoja
        hoja=libro.active
    
        hoja['A1'] = 'Listado de Autos'
   # hoja['B1'] = 'Apellido'
   # hoja['C1'] = 'Salario'
        hoja._add_row
        hoja.append(['Matrícula','Modelo','Anio'])
        objAuto=Auto()
        autos=objAuto.getTodos()
        for fila in autos:
            hoja.append(fila)
        libro.save('ListadoAutos.xlsx')
    
  


        
        ##################################CRUD ACCIDENTE########################
    def accionAccidentes(self):
        self.Accidente=uic.loadUi(r"C:\Users\Zangara\Documents\1-CURSO AX\8460 - TÉCNICAS AVANZADAS DE PROGRAMACIÓN\BBDD\.vscode\LaSegura\CapaVista\CRUDAccidentes.ui")
        self.Accidente.show()

        self.Accidente.btnAgregar.clicked.connect(self.click_btnAgregarAccidente)
        self.Accidente.btnEliminar.clicked.connect(self.click_btnEliminarAccidente)
        self.Accidente.btnBuscar.clicked.connect(self.click_btnBuscarAccidente)
        
        self.Accidente.btnModificar.clicked.connect(self.click_btnModificarAccidente)
        self.Accidente.btnCerrar.clicked.connect(self.click_salirAccidente)

    def click_btnAgregarAccidente(self):
        self.IdAcc=self.Accidente.IdAccidente.text()
        self.Fecha=self.Accidente.FechaAcc.date().toString("yyyy-MM-dd")
        
        self.Lugar=self.Accidente.LugarAcc.toPlainText()

        taccidente=Siniestro()
        taccidente.InsertarInforme(self.IdAcc, self.Fecha, self.Lugar)


    def click_btnEliminarAccidente(self):
        self.IdAcc= self.Accidente.IdAccidente.text()
        tAccidente=Siniestro()
        tAccidente.DeleteInforme(self.IdAcc)
        
    def click_btnBuscarAccidente (self):
        self.IdAcc= self.Accidente.IdAccidente.text()
        tAccidente=Siniestro()
        MostrarAccidente=tAccidente.getInforme(self.IdAcc)
        self.crudautos.FechaAcc.setdate(MostrarAccidente[1])
        self.crudautos.LugarAcc.setText(MostrarAccidente[2])
            
   
    def click_btnModificarAccidente(self):
        self.IdAcc= self.Accidente.IdAccidente.text()
        self.FechaAcc=self.Accidente.FechaAcc.date().toString("yyyy-MM-dd")
        self.Lugar=self.Accidente.LugarAcc.toPlainText()
        tAccidente=Siniestro()
        tAccidente.UpdateInforme(self.IdAcc,self.FechaAcc, self.Lugar )
                
    def click_salirAccidente(self):
        self.Accidente.close()











    def click_btnAgregar(self):
        ######si el texto de una matricula no hay nada, informar error
        # si la matricula tiene el formato correcto, sino informar error
        #si en el texto de la modelo no hay nada informar errror
        #si en el texto del año no hay nada informar error
        #verificar el año sea menor o igual al año actual
        # llamar a la capa negocio para insertar el auto
        self.matricula= self.crudautos.txtMatricula.text()
        self.modelo=self.crudautos.txtModelo.text()
        if self.verificarAnio(self.crudautos.txtAnio.text()):
            self.anio=int(self.crudautos.txtAnio.text())
            if self.verificarMatricula(self.matricula) and self.verificarAnio(self.anio) and self.verificarModelo(self.modelo):

                tautos=Auto()
                tautos.insertAuto(self.matricula, self.modelo,self.anio)
                
        else:
            self.crudautos.lblDisplay.setText('Revise los datos ingresados')


    

    def verificarMatricula(self, matricula):
        patronPatente="[a-zA-Z]{3}\d{3}|[a-zA-Z]{2}\d{3}[a-zA-Z]{2}"
        if matricula =="":
            self.crudautos.lblDisplay.setText('Ingrese una matrícula')
            return False
        elif re.match (patronPatente, matricula)==None:
            self.crudautos.lblDisplay.setText('El formato no concuerda con una patente argentina')
            return False
        else:
            return True
        
    def verificarModelo(self, modelo):
        if modelo =="":
            self.crudautos.lblDisplay.setText('Ingrese un modelo')
            return False
        else:
            return True

    def verificarAnio(self, anio):
        if anio =="":
            self.crudautos.lblDisplay.setText('Ingrese un año')
            return False
        elif (int(anio)<= date.today().year):
            return True
        else:
            return False


    def click_btnEliminar(self):
        self.matricula= self.crudautos.txtMatricula.text()
        if self.verificarMatricula(self.matricula):
            if self.existeMatricula(self.matricula):
                tautos=Auto()
                tautos.borrarAuto(self.matricula)
                self.crudautos.lblDisplay.setText('El vehículo fue eliminado correctamente')
            else:
                self.crudautos.lblDisplay.setText('La matrícula ingresada no existe en la Base de Datos')
        else:
            self.crudautos.lblDisplay.setText('La matrícula no es correcta')

    def existeMatricula (self, matricula):
        tauto=Auto()
        if tauto.getUnAuto(matricula):
            return True
        else:
            return False


    
    def click_btnBuscar(self):
        self.matricula= self.crudautos.txtMatricula.text()
        if self.verificarMatricula(self.matricula):
            if self.existeMatricula(self.matricula):
                tautos=Auto()
                MostrarAuto=tautos.getUnAuto(self.matricula)
                self.crudautos.txtModelo.setText(MostrarAuto[1])
                self.crudautos.txtAnio.setText(MostrarAuto[2])
            else:
                self.crudautos.lblDisplay.setText('La matrícula ingresada no existe en la Base de Datos')
        else:
            self.crudautos.lblDisplay.setText('La matrícula no es correcta')


    def click_btnModificar(self):
        self.matricula= self.crudautos.txtMatricula.text()
        self.Anio=self.crudautos.txtAnio.text()
        self.modelo=self.crudautos.txtModelo.text()
        if self.verificarMatricula(self.matricula):
            if self.existeMatricula(self.matricula):
                tautos=Auto()
                tautos.updateAuto(self.matricula,self.modelo, self.Anio )
                self.crudautos.lblDisplay.setText('Los datos del vehículo han sido actualizados correctamente')
            else:
                self.crudautos.lblDisplay.setText('La matrícula ingresada no existe en la Base de Datos')
        else:
            self.crudautos.lblDisplay.setText('La matrícula no es correcta')

    def click_btnAgregarPersona(self):
        self.carnet= self.crudPersonas.Carnet.text()
        self.nombre=self.crudPersonas.Nombre.text()
        self.direccion=self.crudPersonas.direccion.text()
        persona=Persona()
        persona.insertPersona(self.carnet, self.nombre.upper(), self.direccion.upper())
        self.crudPersonas.lblDisplay.setText('Nuevo Cliente registrado exitosamente')
        self.carnet= self.crudPersonas.Carnet.clear()
        self.nombre=self.crudPersonas.Nombre.clear()
        self.direccion=self.crudPersonas.direccion.clear()


    def click_btnEliminarPersona(self):
        self.id=self.crudPersonas.id.text()
        if self.verificarId(self.id):
            persona=Persona()
            persona.borrarPersona(self.id)
            self.crudPersonas.lblDisplay.setText('Cliente eliminado exitosamente')
        else:
             self.crudautos.lblDisplay.setText('No existe un cliente con ese Id')

    def click_btnBuscarPersona(self):
         self.carnet=self.crudPersonas.Carnet.text()
         if self.VerificarCarnet:
            persona=Persona()
            MostrarPersona=persona.getUnaPersona(self.carnet)
            self.crudPersonas.Nombre.setText(MostrarPersona[2])
            self.crudPersonas.direccion.setText(MostrarPersona[3])
            self.crudPersonas.id.setText(str(MostrarPersona[0]))


    def accionCerrarPersona(self):
        self.crudPersonas.close()

        
    def click_btnModificarPersona(self):
        self.id=self.crudPersonas.id.text()
        if self.verificarId(self.id):
            persona=Persona()
            persona.updatePersona(self.crudPersonas.id.text(),
                                self.crudPersonas.Carnet.text(),
                                self.crudPersonas.Nombre.text(),
                                self.crudPersonas.direccion.text())
            self.crudPersonas.lblDisplay.setText('Datos del Cliente actualizados exitosamente')                       
        else:
            self.crudPersonas.lblDisplay.setText('No existe un cliente con ese Id')

    def VerificarCarnet(self):
        pass
    def verificarId(self, Id):
        persona=Persona()
        if persona.getUnId(Id):
            return True
        else:
            return False

iniciar()

###############AGREGAR BOTON BUSCAR Y ACTUALIZAR
#########AGREGAR EXPRESION REGULAR PARA VALIDAR MATRICULA







# def diasRemanentes (self, tipo):
        
#         self.fecha= datetime(2019,2,28)
        
        
#         pasados=self.fecha + datetime.timedelta(days=tipo)
#         return (datetime.now() - pasados)

# print(diasRemanentes(57))
    

